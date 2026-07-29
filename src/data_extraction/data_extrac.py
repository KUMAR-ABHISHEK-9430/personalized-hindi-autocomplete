# whats in this --> iterate through all the files in docdata and using python-docx extract text 
#  and store that as txt format

# some more things
# data preprocessing will include unicode normalization, remove zwj, zwnj, collapse multiple spaces, collapse balnk lines, 


# flow is through paragraph.text i will get string --> remove multiple spaces -> remove zwj, zwnj --> remove blank lines --> normalize unicode --> store in txt file


import os
import docx
from docx import Document
from pathlib import Path
import re
import unicodedata
import logging



logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
    handlers=[
        logging.FileHandler("preprocessing.log", encoding="utf-8"),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)




class DataCleaner:
    def remove_invisible_characters(self, text):
        # Remove zero-width joiner (ZWJ) and zero-width non-joiner (ZWNJ)
        text = text.replace('\u200d', '')  # ZWJ
        text = text.replace('\u200c', '')  # ZWNJ
        text = text.replace("\ufeff", "")
        return text

    def collapse_multiple_spaces(self, text):
        # Collapse multiple spaces into a single space
        text =  re.sub(r"\n{3,}", "\n\n", text)
        return re.sub(r"[ \t]+", " ", text)

    def normalize_unicode(self, text):
        # Normalize Unicode characters to NFC form
        return unicodedata.normalize('NFC', text)

    def remove_dots(self, text):
        # replaces ...somethign… to something
        text = re.sub(r"[.…]{3,}", " ", text)
        return text

    def replace_phone_numbers(self, text):
        # Replace phone numbers with a placeholder
        phone_pattern = r'(?<!\d)(?:\+91[- ]?|0)?[6-9]\d{9}(?!\d)'
        return re.sub(phone_pattern, '<PHONE>', text)

    def replace_dates(self, text):
        # Replace dates with a placeholder
        date_pattern = r'\b\d{1,2}\s*[./-]\s*\d{1,2}\s*[./-]\s*\d{2,4}\b'
        return re.sub(date_pattern, '<DATE>', text)

    def clean_pipeline(self,text):
        text = self.normalize_unicode(text)
        text = self.remove_invisible_characters(text)
        text = self.replace_phone_numbers(text)
        text = self.replace_dates(text)
        text = self.collapse_multiple_spaces(text)
        text = self.remove_dots(text)
        return text.strip()  # Remove leading/trailing whitespace



# cleaner = DataCleaner()

# cleaned_data = ""



class DataExtractor:
    def __init__(self,dir_path:Path):
        self.cleaner = DataCleaner()
        self.dir_path = Path(dir_path)

    def extract_data(self, doc):
        cleaned_data = ""
        for paragraph in doc.paragraphs:
            if paragraph.text:
                cleaned_data += self.cleaner.clean_pipeline(paragraph.text) + "\n"
        return cleaned_data

    def extract_table_data(self, doc):
        cleaned_data = ""
        for table in doc.tables:
            for row in table.rows:
                cleaned_row_data = [self.cleaner.clean_pipeline(cell.text) for cell in row.cells]
                if len(cleaned_row_data) == 2:  # the second column contains the data we want
                    cleaned_data += cleaned_row_data[1] + "\n"
                elif len(cleaned_row_data) ==  4:  
                    cleaned_data += (cleaned_row_data[3]) + "\n"      
        return cleaned_data

    def process_files(self):
        cleaned_data = []
        for entry in os.scandir(self.dir_path):
            if entry.is_file() and entry.name.endswith('.docx'):
                try:
                    doc = Document(entry.path)
                    cleaned_paragraph_data = self.extract_data(doc)
                    cleaned_table_data = self.extract_table_data(doc)
                    cleaned_data.append(cleaned_paragraph_data)
                    cleaned_data.append(cleaned_table_data)
                except Exception as e:
                    logger.exception(f"Error processing file {entry.name}: {e}")

                logger.info(f"Processed file {entry.name}")
        return "\n".join(cleaned_data)

# with os.scandir(r'C:\projects\auto_complete\personalized-hindi-autocomplete\docdata') as files:
#     for entry in files:
#         if entry.is_file() and entry.name.endswith('.docx'):
#             doc = Document(Path(entry.path))
#             doc = Document(r"C:\projects\auto_complete\personalized-hindi-autocomplete\docdata\_कांड संख्_या-336_25.docx")
#             for index, paragraph in enumerate(doc.paragraphs):
#                 if paragraph.text:
#                     # print(f"Paragraph {index}: {paragraph.text}")
#                     # print("Cleaned Paragraph:", cleaner.clean_pipeline(paragraph.text))
#                     cleaned_data += cleaner.clean_pipeline(paragraph.text) + "\n"
                    
#             for table in doc.tables:
#                 for row in table.rows:
#                     cleaned_row_data = [cleaner.clean_pipeline(cell.text) for cell in row.cells]
#                     cleaned_data += cleaned_row_data[1] + "\n"
#                     # print("Cleaned Row Data:", cleaned_row_data[1])


if __name__ == "__main__":
    dir_path = Path(r'C:\projects\auto_complete\personalized-hindi-autocomplete\docdata')
    extractor = DataExtractor(dir_path)
    cleaned_data = extractor.process_files()
    print("total cleaned data length:", len(cleaned_data))
    with open("cleaned_tokenizer_data.txt", "w", encoding="utf-8") as f:
        f.write(cleaned_data)


